from __future__ import annotations

import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT / "docs" / "meeting-notes" / "教授會議"
OUTPUT = ROOT / "docs" / "meeting-notes" / "教授會議紀錄表_彙整.docx"


@dataclass
class Meeting:
    index: int
    title: str
    date: str
    time: str
    meeting_type: str
    attendees: list[str]
    discussion: list[str]
    decisions: list[str]
    actions: list[str]


def set_cell_text(cell, text: str, bold_prefix: str | None = None, size: float = 10.5) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = "標楷體"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")
    run.font.size = Pt(size)
    if bold_prefix and text.startswith(bold_prefix):
        run.bold = True


def add_runs(paragraph, text: str, *, bold: bool = False, size: float = 10.5) -> None:
    run = paragraph.add_run(text)
    run.font.name = "標楷體"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")
    run.font.size = Pt(size)
    run.bold = bold


def set_paragraph_style(paragraph, *, align=None, before=0, after=0, line=1.15) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def set_table_borders(table, size: str = "8", color: str = "000000") -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_cm: float) -> None:
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm / 2.54 * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def normalize_line(line: str) -> str:
    line = line.strip()
    line = re.sub(r"^\s*[-*]\s+\[ \]\s*", "", line)
    line = re.sub(r"^\s*[-*]\s+", "", line)
    line = re.sub(r"^\s*[•]\s*", "", line)
    line = re.sub(r"\*\*(.*?)\*\*", r"\1", line)
    line = re.sub(r"`([^`]+)`", r"\1", line)
    return line.strip()


def section_lines(content: str, heading: str) -> list[str]:
    pattern = rf"^##\s+{re.escape(heading)}\s*$"
    match = re.search(pattern, content, flags=re.MULTILINE)
    if not match:
        return []
    rest = content[match.end() :]
    next_heading = re.search(r"^##\s+", rest, flags=re.MULTILINE)
    block = rest[: next_heading.start()] if next_heading else rest
    lines = []
    for raw in block.splitlines():
        line = normalize_line(raw)
        if line:
            lines.append(line)
    return lines


def compress_items(lines: list[str]) -> list[str]:
    items: list[str] = []
    for line in lines:
        if re.match(r"^\d+\.\s*$", line):
            continue
        line = re.sub(r"^\d+\.\s+", "", line)
        line = re.sub(r"^\d+[、.]\s*", "", line)
        if line == "-":
            continue
        items.append(line)
    return items


def parse_meeting(path: Path, index: int) -> Meeting:
    content = path.read_text(encoding="utf-8")
    title_match = re.search(r"^#\s+(.+)$", content, flags=re.MULTILINE)
    title = title_match.group(1).strip() if title_match else path.stem
    title = title.replace("與教授開會：", "").replace("與教授開會_", "").replace("與教授開會", "")
    title = title.replace("教授會議：", "").strip("（）()_ ")

    def meta(label: str) -> str:
        m = re.search(rf"^-\s*{label}:\s*(.+)$", content, flags=re.MULTILINE)
        return m.group(1).strip() if m else ""

    attendees_raw = meta("與會者")
    attendees = [x.strip() for x in re.split(r"[,，、]", attendees_raw) if x.strip()]
    return Meeting(
        index=index,
        title=title or path.stem,
        date=meta("會議日期"),
        time=meta("會議時間"),
        meeting_type=meta("會議類型") or "教授會議",
        attendees=attendees,
        discussion=compress_items(section_lines(content, "討論內容")),
        decisions=compress_items(section_lines(content, "決議事項")),
        actions=compress_items(section_lines(content, "後續行動項目")),
    )


def roc_date(iso_date: str) -> str:
    if not iso_date:
        return "未記錄"
    dt = datetime.strptime(iso_date, "%Y-%m-%d")
    return f"中華民國 {dt.year - 1911} 年 {dt.month:02d} 月 {dt.day:02d} 日"


def next_date_text(meetings: list[Meeting], idx: int) -> str:
    if idx + 1 >= len(meetings):
        return "未定"
    return roc_date(meetings[idx + 1].date)


def split_time(meeting: Meeting) -> tuple[str, str]:
    if not meeting.time:
        return "未記錄", "未記錄"
    text = meeting.time.replace("－", "-").replace("–", "-").replace("~", "-")
    if "-" in text:
        start, end = [part.strip() for part in text.split("-", 1)]
        return start, end
    if "約" in text and "分鐘" in text:
        return "未記錄", text
    return text, "未記錄"


def attendee_block(attendees: list[str]) -> str:
    professor = [a.replace("(教授)", "").replace("（教授）", "").strip() for a in attendees if "教授" in a]
    members = [a for a in attendees if "教授" not in a]
    lines = []
    if professor:
        lines.append("指導老師　" + "、".join(professor))
    if members:
        lines.append("組員　" + "、".join(members))
    return "\n".join(lines) if lines else "未記錄"


def add_numbered_lines(
    cell,
    lines: list[str],
    empty_text="未記錄",
    max_items: int | None = None,
    *,
    clear: bool = True,
) -> None:
    if clear:
        cell.text = ""
    if not lines:
        p = cell.paragraphs[0] if clear else cell.add_paragraph()
        set_paragraph_style(p, after=0)
        add_runs(p, empty_text)
        return
    selected = lines[:max_items] if max_items else lines
    for i, item in enumerate(selected, 1):
        p = cell.paragraphs[0] if clear and i == 1 else cell.add_paragraph()
        set_paragraph_style(p, after=2, line=1.12)
        p.paragraph_format.first_line_indent = Cm(-0.6)
        p.paragraph_format.left_indent = Cm(0.6)
        add_runs(p, f"{i}、", size=10)
        add_runs(p, item, size=10)
    if max_items and len(lines) > max_items:
        p = cell.add_paragraph()
        set_paragraph_style(p, after=2, line=1.12)
        p.paragraph_format.first_line_indent = Cm(-0.6)
        p.paragraph_format.left_indent = Cm(0.6)
        add_runs(p, f"{max_items + 1}、", size=10)
        add_runs(p, f"其餘細節請參考原始會議紀錄（共 {len(lines)} 項）。", size=10)


def add_content_cell(cell, meeting: Meeting) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph_style(p, after=2)
    add_runs(p, "會議內容：", bold=True, size=10.5)

    def block(title: str, lines: list[str], limit: int | None = None) -> None:
        heading = cell.add_paragraph()
        set_paragraph_style(heading, before=2, after=1)
        add_runs(heading, title, bold=True, size=10.5)
        if not lines:
            body = cell.add_paragraph()
            set_paragraph_style(body, after=2)
            add_runs(body, "未記錄", size=10)
            return
        selected = lines[:limit] if limit else lines
        for i, item in enumerate(selected, 1):
            body = cell.add_paragraph()
            set_paragraph_style(body, after=2, line=1.1)
            body.paragraph_format.first_line_indent = Cm(-0.6)
            body.paragraph_format.left_indent = Cm(0.6)
            add_runs(body, f"{i}、", size=10)
            add_runs(body, item, size=10)
        if limit and len(lines) > limit:
            body = cell.add_paragraph()
            set_paragraph_style(body, after=2, line=1.1)
            body.paragraph_format.first_line_indent = Cm(-0.6)
            body.paragraph_format.left_indent = Cm(0.6)
            add_runs(body, f"{limit + 1}、", size=10)
            add_runs(body, f"其餘細節請參考原始會議紀錄（共 {len(lines)} 項）。", size=10)

    block("一、討論內容", meeting.discussion)
    block("二、決議事項", meeting.decisions)


def add_header_title(doc: Document, number: int) -> None:
    p = doc.add_paragraph()
    set_paragraph_style(p, align=WD_ALIGN_PARAGRAPH.CENTER, after=0, line=1.0)
    add_runs(p, "國立臺北商業大學　資訊管理系", bold=True, size=18)
    p = doc.add_paragraph()
    set_paragraph_style(p, align=WD_ALIGN_PARAGRAPH.CENTER, after=4, line=1.0)
    add_runs(p, f"第{number}次專題討論　會議紀錄表", bold=True, size=18)


def add_meeting_table(doc: Document, meeting: Meeting, meetings: list[Meeting], idx: int) -> None:
    add_header_title(doc, meeting.index)
    start, end = split_time(meeting)

    table = doc.add_table(rows=8, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, size="8")
    widths = [4.3, 5.0, 3.8, 5.1]
    for row in table.rows:
        prevent_row_split(row)
        for i, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            set_cell_width(cell, widths[i])

    cells = table.rows[0].cells
    cells[0].merge(cells[1])
    cells[2].merge(cells[3])
    set_cell_text(cells[0], "會議日期：" + roc_date(meeting.date), "會議日期：")
    set_cell_text(cells[2], "會議地點：未記錄", "會議地點：")

    cells = table.rows[1].cells
    cells[0].merge(cells[1])
    cells[2].merge(cells[3])
    chair = "蒯思齊" if any("教授" in a for a in meeting.attendees) else "未記錄"
    set_cell_text(cells[0], f"會議主席：{chair}", "會議主席：")
    set_cell_text(cells[2], "會議記錄：未記錄", "會議記錄：")

    cells = table.rows[2].cells
    cells[0].merge(cells[1])
    cells[2].merge(cells[3])
    set_cell_text(cells[0], f"開會時間：{start}", "開會時間：")
    set_cell_text(cells[2], f"散會時間：{end}", "散會時間：")

    topic = f"會議議題：{meeting.title or meeting.meeting_type}"
    row = table.rows[3]
    row.cells[0].merge(row.cells[3])
    set_cell_text(row.cells[0], topic, "會議議題：")

    row = table.rows[4]
    row.cells[0].merge(row.cells[3])
    set_cell_text(row.cells[0], "開會人員：\n" + attendee_block(meeting.attendees), "開會人員：")

    row = table.rows[5]
    row.cells[0].merge(row.cells[3])
    set_cell_text(row.cells[0], "缺席人員：無", "缺席人員：")

    row = table.rows[6]
    row.cells[0].merge(row.cells[3])
    add_content_cell(row.cells[0], meeting)

    row = table.rows[7]
    row.cells[0].merge(row.cells[3])
    cell = row.cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph_style(p, after=2)
    add_runs(p, "下次開會時間：" + next_date_text(meetings, idx), bold=True, size=10.5)
    p = cell.add_paragraph()
    set_paragraph_style(p, before=4, after=2)
    add_runs(p, "下次會議內容：", bold=True, size=10.5)
    add_numbered_lines(cell, meeting.actions, empty_text="未記錄", clear=False)


def build() -> None:
    paths = sorted(SOURCE_DIR.glob("*.md"))
    meetings = [parse_meeting(path, i + 1) for i, path in enumerate(paths)]

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "標楷體"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")
    normal.font.size = Pt(10.5)

    for idx, meeting in enumerate(meetings):
        if idx:
            doc.add_section(WD_SECTION.NEW_PAGE)
        add_meeting_table(doc, meeting, meetings, idx)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
